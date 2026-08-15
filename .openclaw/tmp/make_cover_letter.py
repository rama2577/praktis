#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a clean cover letter .docx for the Sony Music Finance Director role."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = "/Users/staff/.openclaw-autoclaw/workspace/Sony_Music_Finance_Director_Cover_Letter_Rama_Wijaya.docx"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

# Page margins
for s in doc.sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)


def p(text="", bold=False, size=11, space_after=8, align=None, color=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        par.alignment = align
    if text:
        run = par.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return par


# --- Header block (name + contact) ---
p("Rama Wijaya", bold=True, size=16, space_after=2)
p("Bandung / Jakarta, Indonesia", size=10.5, space_after=1)
p("Phone: [your phone]  •  Email: [your email]", size=10.5, space_after=1)
p("LinkedIn: linkedin.com/in/rama-wijaya-supplychain", size=10.5, space_after=10)

# --- Date ---
p("[Date]", size=11, space_after=10)

# --- Recipient ---
p("Hiring Manager", space_after=1)
p("Sony Music Entertainment", space_after=1)
p("Indonesia", space_after=10)

# --- Subject line ---
p("Re: Application for Finance Director, Indonesia", bold=True, size=11.5, space_after=10)

# --- Salutation ---
p("Dear Hiring Manager,", space_after=8)

# --- Body ---
body = [
    ("I am writing to express my strong interest in the Finance Director position at Sony Music "
     "Entertainment Indonesia. The opportunity to bring commercial finance leadership to the creative "
     "energy of the music industry — and to help shape the strategic direction of Sony Music's Southeast "
     "Asia business — is precisely the kind of mandate I have built my career around."),

    ("Over more than 20 years in senior commercial and operational leadership roles, I have consistently "
     "operated at the intersection of strategy, finance, and business growth. As Vice President of Business "
     "Development & Partnership at Mile.app, I lead commercial deal evaluation, financial modelling, and "
     "business-case development for partnerships that directly drive revenue — work that closely mirrors the "
     "commercial transaction and artist-investment analysis this role requires. I own P&L accountability, "
     "build forecasts, and translate financial insight into recommendations that balance commercial ambition "
     "with disciplined risk management."),

    ("My background spans both sides of this mandate — the commercial and the controllership. On financial "
     "control and operational excellence, I have delivered measurable results: I led an operational turnaround "
     "that lifted delivery performance from 50% to 97%, built a last-mile platform processing more than one "
     "million shipments per day, and drove a 20% reduction in operating costs through process improvement and "
     "automation. As a Six Sigma Black Belt, I have spent my career strengthening governance, internal controls, "
     "and cost discipline — the same rigour required to safeguard assets, manage cash flow and liquidity, and "
     "ensure IFRS, tax, and statutory compliance."),

    ("I hold a Bachelor's degree in Fiscal Administration from the University of Indonesia, which gives me a "
     "strong foundation in tax, financial reporting, and regulatory compliance, complemented by an MASc in "
     "Logistics (MITx) and advanced financial-modelling capability. I am fluent in English and a native Bahasa "
     "Indonesia speaker, and I have spent years partnering with regional leadership and cross-functional teams "
     "within a matrix, multinational environment."),

    ("What excites me most about Sony Music is the chance to bring this blend of commercial acumen and financial "
     "control to a creative industry I care about deeply. I would welcome the opportunity to discuss how I can "
     "help the country team evaluate artist signings and licensing deals with confidence, drive performance, and "
     "build a high-performing finance function."),
]

for para in body:
    p(para, size=11, space_after=8)

# --- Closing ---
p("Thank you for your consideration. I look forward to the opportunity to discuss my application further.",
  space_after=14)

p("Sincerely,", space_after=2)
p("Rama Wijaya", space_after=0)

doc.save(OUT)
print("Saved:", OUT)
